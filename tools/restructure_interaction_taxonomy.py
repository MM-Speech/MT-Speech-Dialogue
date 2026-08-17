from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\26547\Desktop\read\readme.docx")
OUTPUT = SOURCE.with_name("readme_restructured.docx")

HEADINGS = {
    "evaluation": {
        "turn": "Evaluation Focus: Turn-taking & Interruption (轮次衔接与打断)",
        "full": "Evaluation Focus: Full-duplex Interaction (全双工交互)",
        "other": "Evaluation Focus: Other Dialogue Capabilities (其他对话能力)",
    },
    "mode": {
        "turn": "Interaction Mode: Turn-taking (轮流对话)",
        "half": "Interaction Mode: Half-duplex / Barge-in (半双工 / 带打断)",
        "full": "Interaction Mode: Full-duplex (全双工)",
    },
}

EVALUATION = {
    "turn": {
        "Talking Turns: Benchmarking Audio Foundation Models on Turn-Taking Dynamics",
        "Investigating the Impact of Incremental Processing and Voice Activity Projection on Spoken Dialogue Systems",
    },
    "full": {
        "Full-Duplex-Bench: A Benchmark to Evaluate Full-duplex Spoken Dialogue Models on Turn-taking Capabilities",
        "The ICASSP 2026 HumDial Challenge: Benchmarking Human-like Spoken Dialogue Systems in the LLM Era",
    },
    "other": {
        "SD-Eval: A Benchmark Dataset for Spoken Dialogue Understanding Beyond Words",
        "Contextual Interactive Evaluation of TTS Models in Dialogue Systems",
        "URO-Bench: Towards Comprehensive Evaluation for End-to-End Spoken Dialogue Models",
        "WavReward: Spoken Dialogue Models With Generalist Reward Evaluators",
        "MAD: A Benchmark for Multi-Turn Audio Dialogue Fact-Checking",
        "VoxRole: A Comprehensive Benchmark for Evaluating Speech-Based Role-Playing Agents",
        "Evaluating Bias in Spoken Dialogue LLMs for Real-World Decisions and Recommendations (FairDialogue)",
        "MULTI-Bench: A Multi-turn Interactive Benchmark for Assessing Emotional Intelligence Ability of Spoken Dialogue Models",
        "Audio MultiChallenge: A Multi-Turn Evaluation of Spoken Dialogue Systems on Natural Human Interaction",
        "VoxDialogue: Can Spoken Dialogue Systems Understand Information Beyond Words?",
        "EchoMind: An Interrelated Multi-level Benchmark for Evaluating Empathetic Speech Language Models",
        "HumDial-EIBench: A Human-Recorded Multi-Turn Emotional Intelligence Benchmark for Audio Language Models",
        "VoiceBench: Benchmarking LLM-Based Voice Assistants",
    },
}

MODES = {
    "Models": {
        "turn": {
            "SpeechGPT: Empowering Large Language Models with Intrinsic Cross-Modal Conversational Abilities",
            "Qwen2-Audio Technical Report", "Style-Talker: Finetuning Audio Language Model and Style-Based Text-to-Speech Model for Fast Spoken Dialogue Generation",
            "EMOVA: Empowering Language Models to See, Hear and Speak with Vivid Emotions",
            "Internalizing ASR with Implicit Chain of Thought for Efficient Speech-to-Speech Conversational LLM",
            "Building a Taiwanese Mandarin Spoken Language Model: A First Attempt",
            "GLM-4-Voice: Towards Intelligent and Human-Like End-to-End Spoken Chatbot",
            "OpenOmni: Advancing Open-Source Omnimodal Large Language Models with Progressive Multimodal Alignment and Real-time Emotional Speech Synthesis",
            "Qwen2.5-Omni Technical Report", "LLaMA-Omni 2: LLM-based Real-time Spoken Chatbot with Autoregressive Streaming Speech Synthesis",
            "GOAT-SLM: A Spoken Language Model with Paralinguistic and Speaker Characteristic Awareness",
            "OpenS2S: Advancing Fully Open-Source End-to-End Empathetic Large Speech Language Model",
            "Step-Audio 2 Technical Report", "InteractiveOmni: A Unified Omni-modal Model for Audio-Visual Multi-turn Dialogue",
        },
        "half": {
            "Mini-Omni: Language Models Can Hear, Talk While Thinking in Streaming",
            "Freeze-Omni: A Smart and Low Latency Speech-to-Speech Dialogue Model with Frozen LLM",
        },
        "full": {
            "Beyond Turn-Based Interfaces: Synchronous LLMs as Full-Duplex Dialogue Agents",
            "Moshi: A Speech-Text Foundation Model for Real-Time Dialogue",
            "SALMONN-omni: A Codec-free LLM for Full-duplex Speech Understanding and Generation",
            "MinMo: A Multimodal Large Language Model for Seamless Voice Interaction",
            "SALM-Duplex: Efficient and Direct Duplex Modeling for Speech-to-Speech Language Model",
            "OmniFlatten: An End-to-end GPT Model for Seamless Voice Conversation",
            "FlashLabs Chroma 1.0: A Real-Time End-to-End Spoken Dialogue Model with Personalized Voice Cloning",
            "PersonaPlex: Voice and Role Control for Full Duplex Conversational Speech Models",
        },
    },
    "Training Methods": {
        "turn": {
            "SpeechAlign: Aligning Speech Generation to Human Preferences",
            "WavRAG: Audio-Integrated Retrieval Augmented Generation for Spoken Dialogue Models",
            "Enhancing Speech-to-Speech Dialogue Modeling with End-to-End Retrieval-Augmented Generation",
            "Aligning Spoken Dialogue Models from User Interactions",
            "Chain-of-Thought Training for Open E2E Spoken Dialogue Systems",
            "Stream RAG: Instant and Accurate Spoken Dialogue Systems with Streaming Tool Usage",
            "VoiceTextBlender: Augmenting Large Language Models with Speech Capabilities via Single-Stage Joint Speech-Text Supervised Fine-Tuning",
            "Optimizing Conversational Quality in Spoken Dialogue Systems with Reinforcement Learning from AI Feedback",
        },
        "full": {
            "SHANKS: Simultaneous Hearing and Thinking for Spoken Language Models",
            "Multi-Faceted Interactivity Alignment in Full-Duplex Speech Models",
        },
    },
}


def title(paragraph):
    text = paragraph.text.strip()
    return text.split('". ', 1)[0][1:] if text.startswith('"') else None


def papers_in_section(document, section_name):
    active, result = False, []
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 1":
            if active:
                break
            active = paragraph.text.strip() == section_name
        elif active and title(paragraph):
            result.append(paragraph)
    return result


def new_num(document, source_id):
    numbering = document.part.numbering_part.element
    source = next(n for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId")) == source_id)
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    clone = deepcopy(source)
    clone.set(qn("w:numId"), str(max(ids) + 1))
    numbering.append(clone)
    return clone.get(qn("w:numId"))


def set_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    for node in list(num_pr):
        if node.tag in {qn("w:ilvl"), qn("w:numId")}:
            num_pr.remove(node)
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId"); num.set(qn("w:val"), num_id)
    num_pr.append(ilvl); num_pr.append(num)


def insert_heading(document, cursor, label):
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.add_run(label)
    cursor.addnext(paragraph._p)
    return paragraph._p


def restructure(document, section_name, groups, heading_labels):
    heading = next(p for p in document.paragraphs if p.style.name == "Heading 1" and p.text.strip() == section_name)
    papers = papers_in_section(document, section_name)
    by_title = {title(p): p for p in papers}
    assigned = set().union(*groups.values())
    if assigned != set(by_title):
        raise ValueError(f"{section_name}: missing={set(by_title) - assigned}; unknown={assigned - set(by_title)}")

    cursor = heading._p
    for group in ("turn", "half", "full", "other"):
        titles = groups.get(group, set())
        if not titles:
            continue
        cursor = insert_heading(document, cursor, heading_labels[group])
        first = by_title[next(iter(titles))]
        old_id = first._p.pPr.find(qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
        new_id = new_num(document, old_id)
        for paragraph in papers:
            if title(paragraph) in titles:
                cursor.addnext(paragraph._p)
                set_num(paragraph, new_id)
                cursor = paragraph._p


def main():
    document = Document(SOURCE)
    restructure(document, "Evaluation & Benchmarks", EVALUATION, HEADINGS["evaluation"])
    for section, groups in MODES.items():
        restructure(document, section, groups, HEADINGS["mode"])
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
