from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "README.md"
README_OUTPUT = ROOT / "README1.md"
DOCX_OUTPUT = ROOT / "readme1.docx"

SECTION_ORDER = [
    "Survey",
    "General Surveys",
    "Datasets",
    "Evaluation & Benchmarks",
    "Models",
    "Training Methods",
]

SECTION_LABELS = {
    "Survey": "Survey (\u591a\u8f6e\u8bed\u97f3\u5bf9\u8bdd\u7efc\u8ff0\u7c7b)",
    "General Surveys": "General Surveys (\u901a\u7528\u7efc\u8ff0)",
    "Datasets": "Datasets (\u6570\u636e\u96c6\u7c7b)",
    "Evaluation & Benchmarks": "Evaluation & Benchmarks (\u8bc4\u6d4b\u7c7b)",
    "Models": "Models (\u6a21\u578b\u7c7b)",
    "Training Methods": "Training Methods (\u8bad\u7ec3\u65b9\u6cd5\u7c7b)",
}

GENERAL_SURVEYS = {
    "A Review of Dialogue Systems: From Trained Monkeys to Stochastic Parrots",
    "Transformers in Speech Processing: A Survey",
    "A Survey on Speech Large Language Models for Understanding",
    "Preference Tuning with Human Feedback on Language, Speech, and Vision Tasks: A Survey",
    "Aligning Multimodal LLM with Human Preference: A Survey",
}

EXCLUDE = {
    "ComperDial: Commonsense Persona-grounded Dialogue Dataset and Benchmark",
    "MARS-Bench: A Multi-turn Athletic Real-world Scenario Benchmark for Dialogue Evaluation",
    "AudioGPT: Understanding and Generating Speech, Music, Sound, and Talking Head",
    "AudioPaLM: A Large Language Model That Can Speak and Listen",
    "Dialog Action-Aware Transformer for Dialog Policy Learning",
    "Conversation Forests: The Key to Fine Tuning Large Language Models for Multi-Turn Medical Conversations is Branching",
}

MOVE_TO = {
    "SD-Eval: A Benchmark Dataset for Spoken Dialogue Understanding Beyond Words": "Evaluation & Benchmarks",
    "MAD: A Benchmark for Multi-Turn Audio Dialogue Fact-Checking": "Evaluation & Benchmarks",
    "Audio MultiChallenge: A Multi-Turn Evaluation of Spoken Dialogue Systems on Natural Human Interaction": "Evaluation & Benchmarks",
    "The ICASSP 2026 HumDial Challenge: Benchmarking Human-like Spoken Dialogue Systems in the LLM Era": "Evaluation & Benchmarks",
    "InteractiveOmni: A Unified Omni-modal Model for Audio-Visual Multi-turn Dialogue": "Models",
    "OpenOmni: Advancing Open-Source Omnimodal Large Language Models with Progressive Multimodal Alignment and Real-time Emotional Speech Synthesis": "Models",
    "WavRAG: Audio-Integrated Retrieval Augmented Generation for Spoken Dialogue Models": "Training Methods",
    "Stream RAG: Instant and Accurate Spoken Dialogue Systems with Streaming Tool Usage": "Training Methods",
    "SHANKS: Simultaneous Hearing and Thinking for Spoken Language Models": "Training Methods",
}

PAPER_PATTERN = re.compile(
    r'^\d+\. \*\*"(?P<title>.+?)"\*\*\. \*(?P<authors>.+?)\* (?P<venue>.+?)\. \[\[Paper\]\((?P<url>[^)]+)\)\]$'
)


def chronological_key(paper):
    """Sort by the displayed publication year, then by arXiv submission month."""
    years = re.findall(r"20\d{2}", paper["venue"])
    year = int(years[-1]) if years else 9999
    arxiv = re.search(r"arxiv\.org/abs/(\d{2})(\d{2})\.\d+", paper["url"])
    month = int(arxiv.group(2)) if arxiv else 13
    return year, month, paper["title"].lower()


def read_papers():
    papers = {section: [] for section in SECTION_ORDER}
    section = None
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            section = line[3:]
            continue
        match = PAPER_PATTERN.match(line)
        if not match or section is None:
            continue
        paper = match.groupdict()
        title = paper["title"]
        if title in EXCLUDE:
            continue
        target = "General Surveys" if title in GENERAL_SURVEYS else MOVE_TO.get(title, section)
        papers[target].append(paper)
    for entries in papers.values():
        entries.sort(key=chronological_key)
    return papers


def markdown(papers):
    lines = [
        "# MT-Speech-Dialogue",
        "",
        "> A curated paper list for **Multi-Turn Speech Dialogue**.",
        ">",
        "> Scope: papers directly addressing multi-turn, end-to-end, or full-duplex speech interaction. Broad surveys are listed separately under **General Surveys**.",
    ]
    for section in SECTION_ORDER:
        lines.extend(["", f"## {SECTION_LABELS[section]}", ""])
        for number, paper in enumerate(papers[section], start=1):
            lines.append(
                f'{number}. **"{paper["title"]}"**. *{paper["authors"]}* {paper["venue"]}. [[Paper]({paper["url"]})]'
            )
    return "\n".join(lines) + "\n"


def append_hyperlink(paragraph, text, url):
    hyperlink = OxmlElement("w:hyperlink")
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def document(papers):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name, size in (("Heading 1", 18), ("Heading 2", 14)):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)

    doc.add_heading("Multi-Turn Speech Dialogue", level=0)
    scope = doc.add_paragraph()
    scope.add_run("Scope: ").bold = True
    scope.add_run(
        "papers directly addressing multi-turn, end-to-end, or full-duplex speech interaction. "
        "Broad surveys are listed separately under General Surveys."
    )

    for section in SECTION_ORDER:
        doc.add_heading(SECTION_LABELS[section], level=1)
        for number, paper in enumerate(papers[section], start=1):
            paragraph = doc.add_paragraph(style="Normal")
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.add_run(f'{number}. "{paper["title"]}". ')
            authors = paragraph.add_run(f'{paper["authors"]} ')
            authors.italic = True
            paragraph.add_run(f'{paper["venue"]}. ')
            append_hyperlink(paragraph, "Paper", paper["url"])
    doc.save(DOCX_OUTPUT)


def main():
    papers = read_papers()
    expected_count = 65
    total = sum(len(entries) for entries in papers.values())
    if total != expected_count:
        raise ValueError(f"Expected {expected_count} retained papers, found {total}")
    README_OUTPUT.write_text(markdown(papers), encoding="utf-8")
    document(papers)
    for section in SECTION_ORDER:
        print(f"{section}: {len(papers[section])}")
    print(f"Total: {total}")


if __name__ == "__main__":
    main()
