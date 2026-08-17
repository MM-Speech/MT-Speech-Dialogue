from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


WORKSPACE = Path(r"D:\vscode\MT-Speech-Dialogue")
BASE_DOCX = Path(r"C:\Users\26547\Desktop\read\readme.docx")
OUTPUT_DOCX = WORKSPACE / "readme_pre_2023.docx"
MIN_YEAR = 0
DEDUPLICATE_AGAINST_BASE = True
DOCUMENT_TITLE = "Multi-Turn Speech Dialogue: Pre-2023 Supplement"
INTRO_TEXT = "A curated supplement of papers published in 2022 or earlier, excluding titles already listed in readme.docx."
CUTOFF_TEXT = "2022 and earlier"
COUNT_LABEL = "New papers"
SUMMARY_THIRD_LABEL = "Deduplication"
SUMMARY_THIRD_TEXT = None
EMPTY_SECTION_TEXT = "No additional eligible pre-2023 paper was found after deduplication. Early work in this area is represented under Half-duplex / Controlled Barge-in rather than modern native full-duplex interaction."
HEADER_TEXT = "Multi-Turn Speech Dialogue | Pre-2023 Supplement"
SUBJECT_TEXT = "Curated papers published in 2022 or earlier"
RECLASSIFICATIONS = {}


def paper(title, authors, venue, year, url):
    return {"title": title, "authors": authors, "venue": venue, "year": year, "url": url}


SECTIONS = [
    ("Survey", None, [
        paper("Spoken Dialogue Technology: Enabling the Conversational User Interface", "Michael F. McTear", "ACM Computing Surveys", 2002, "https://doi.org/10.1145/505282.505285"),
        paper("POMDP-Based Statistical Spoken Dialog Systems: A Review", "Steve Young et al.", "Proceedings of the IEEE", 2013, "https://doi.org/10.1109/JPROC.2012.2225812"),
        paper("Using Neural Networks for Data-Driven Backchannel Prediction: A Survey on Input Features and Training Techniques", "Markus Mueller et al.", "HCI International", 2015, "https://doi.org/10.1007/978-3-319-20916-6_31"),
        paper("A Survey of Available Corpora For Building Data-Driven Dialogue Systems: The Journal Version", "Iulian V. Serban et al.", "Dialogue & Discourse", 2018, "https://aclanthology.org/2018.dnd-9.7/"),
        paper("Turn-taking in Conversational Systems and Human-Robot Interaction: A Review", "Gabriel Skantze", "Computer Speech & Language", 2021, "https://doi.org/10.1016/j.csl.2020.101178"),
    ]),
    ("General Surveys", None, [
        paper("A Survey on Dialogue Systems: Recent Advances and New Frontiers", "Hongshen Chen et al.", "SIGKDD Explorations", 2017, "https://arxiv.org/abs/1711.01731"),
        paper("Neural Approaches to Conversational AI", "Jianfeng Gao et al.", "Foundations and Trends in Information Retrieval", 2019, "https://arxiv.org/abs/1809.08267"),
        paper("Recent Neural Methods on Dialogue State Tracking for Task-Oriented Dialogue Systems: A Survey", "Vevake Balaraman et al.", "SIGDIAL", 2021, "https://aclanthology.org/2021.sigdial-1.25/"),
        paper("Who Says What to Whom: A Survey of Multi-Party Conversations", "Jia-Chen Gu et al.", "IJCAI", 2022, "https://www.ijcai.org/proceedings/2022/768"),
    ]),
    ("Datasets", None, [
        paper("The HCRC Map Task Corpus", "Anne H. Anderson et al.", "Language and Speech", 1991, "https://doi.org/10.1177/002383099103400404"),
        paper("SWITCHBOARD: Telephone Speech Corpus for Research and Development", "John J. Godfrey et al.", "ICASSP", 1992, "https://doi.org/10.1109/ICASSP.1992.225858"),
        paper("The ICSI Meeting Corpus", "Adam Janin et al.", "ICASSP", 2003, "https://doi.org/10.1109/ICASSP.2003.1198793"),
        paper("The Fisher Corpus: A Resource for the Next Generations of Speech-to-Text", "Christopher Cieri et al.", "LREC", 2004, "https://aclanthology.org/L04-1500/"),
        paper("The AMI Meeting Corpus", "Wessel Kraaij et al.", "MLMI", 2005, "https://groups.inf.ed.ac.uk/ami/corpus/"),
        paper("HKUST/MTS: A Very Large Scale Mandarin Telephone Speech Corpus", "Yi Liu et al.", "ISCSLP", 2006, "https://doi.org/10.1007/11939993_73"),
        paper("IEMOCAP: Interactive Emotional Dyadic Motion Capture Database", "Carlos Busso et al.", "Language Resources and Evaluation", 2008, "https://doi.org/10.1007/s10579-008-9076-6"),
        paper("The SEMAINE Database: Annotated Multimodal Records of Emotionally Colored Conversations between a Person and a Limited Agent", "Gary McKeown et al.", "IEEE Transactions on Affective Computing", 2012, "https://doi.org/10.1109/T-AFFC.2011.25"),
        paper("The MAHNOB Mimicry Database: A Database of Naturalistic Human Interactions", "Sanjay Bilakhia et al.", "Pattern Recognition Letters", 2015, "https://doi.org/10.1016/j.patrec.2015.03.005"),
        paper("The NoXi Database: Multimodal Recordings of Mediated Novice-Expert Interactions", "Angelo Cafaro et al.", "ICMI", 2017, "https://doi.org/10.1145/3136755.3136780"),
        paper("Japanese Dialogue Corpus of Information Navigation and Attentive Listening Annotated with Extended ISO-24617-2 Dialogue Act Tags", "Koichiro Yoshino et al.", "LREC", 2018, "https://aclanthology.org/L18-1462/"),
        paper("Taskmaster-1: Toward a Realistic and Diverse Dialog Dataset", "Bill Byrne et al.", "EMNLP-IJCNLP", 2019, "https://aclanthology.org/D19-1459/"),
        paper("Alexa in the Wild: Collecting Unconstrained Conversations with a Modern Voice Assistant in a Public Environment", "Benjamin Benk et al.", "LREC", 2020, "https://aclanthology.org/2020.lrec-1.77/"),
        paper("HarperValleyBank: A Domain-Specific Spoken Dialog Corpus", "Mike Wu et al.", "arXiv", 2020, "https://arxiv.org/abs/2010.13929"),
        paper("The MSP-Conversation Corpus", "Luz Martinez-Lucas et al.", "Interspeech", 2020, "https://www.isca-archive.org/interspeech_2020/martinezlucas20_interspeech.html"),
        paper("ASCEND: A Spontaneous Chinese-English Dataset for Code-switching in Multi-turn Conversation", "Holy Lovenia et al.", "arXiv", 2021, "https://arxiv.org/abs/2112.06223"),
        paper("Collection and Analysis of Travel Agency Task Dialogues with Age-Diverse Speakers", "Michimasa Inaba et al.", "LREC", 2022, "https://aclanthology.org/2022.lrec-1.619/"),
        paper("Design and Evaluation of the Corpus of Everyday Japanese Conversation", "Hanae Koiso et al.", "LREC", 2022, "https://aclanthology.org/2022.lrec-1.599/"),
    ]),
    ("Evaluation & Benchmarks", "Evaluation Focus: Turn-taking & Interruption", [
        paper("Effects of System Barge-In Responses on User Impressions", "Jun-ichi Hirasawa et al.", "Eurospeech", 1999, "https://www.isca-archive.org/eurospeech_1999/hirasawa99_eurospeech.html"),
        paper("Learning Decision Trees to Determine Turn-Taking by Spoken Dialogue Systems", "Ryo Sato et al.", "ICSLP", 2002, "https://www.isca-archive.org/icslp_2002/sato02b_icslp.html"),
        paper("Multimodal End-of-Turn Prediction in Multi-Party Meetings", "Iwan de Kok and Dirk Heylen", "ICMI-MLMI", 2009, "https://doi.org/10.1145/1647314.1647332"),
        paper("A Finite-State Turn-Taking Model for Spoken Dialog Systems", "Antoine Raux and Maxine Eskenazi", "NAACL-HLT", 2009, "https://aclanthology.org/N09-1071/"),
        paper("Turn-Taking Cues in Task-Oriented Dialogue", "Agustin Gravano and Julia Hirschberg", "Computer Speech & Language", 2011, "https://doi.org/10.1016/j.csl.2010.10.003"),
        paper("Evaluation and Optimisation of Incremental Processors", "Okko Buss and David Schlangen", "Dialogue & Discourse", 2011, "https://aclanthology.org/2011.dnd-2.10/"),
        paper("Continuously Predicting and Processing Barge-in During a Live Spoken Dialogue Task", "Ethan O. Selfridge et al.", "SIGDIAL", 2013, "https://aclanthology.org/W13-4063/"),
        paper("Predicting User Satisfaction from Turn-Taking in Spoken Conversations", "Shammur Absar Chowdhury et al.", "Interspeech", 2016, "https://www.isca-archive.org/interspeech_2016/chowdhury16_interspeech.html"),
        paper("Towards a General, Continuous Model of Turn-Taking in Spoken Dialogue Using LSTM Recurrent Neural Networks", "Gabriel Skantze", "SIGDIAL", 2017, "https://aclanthology.org/W17-5527/"),
        paper("Towards Deep End-of-Turn Prediction for Situated Spoken Dialogue Systems", "Angelika Maier et al.", "Interspeech", 2017, "https://www.isca-archive.org/interspeech_2017/maier17_interspeech.html"),
        paper("Turn-Taking Estimation Model Based on Joint Embedding of Lexical and Prosodic Contents", "Chaoran Liu et al.", "Interspeech", 2017, "https://www.isca-archive.org/interspeech_2017/liu17_interspeech.html"),
        paper("Improving End-of-Turn Detection in Spoken Dialogues by Detecting Speaker Intentions as a Secondary Task", "Zakaria Aldeneh et al.", "ICASSP", 2018, "https://doi.org/10.1109/ICASSP.2018.8461340"),
        paper("Investigating Speech Features for Continuous Turn-Taking Prediction Using LSTMs", "Matthew Roddy et al.", "Interspeech", 2018, "https://www.isca-archive.org/interspeech_2018/roddy18_interspeech.html"),
        paper("Neural Dialogue Context Online End-of-Turn Detection", "Ryo Masumura et al.", "SIGDIAL", 2018, "https://aclanthology.org/W18-5024/"),
        paper("Prediction of Turn-Taking Using Multitask Learning with Prediction of Backchannels and Fillers", "Kohei Hara et al.", "Interspeech", 2018, "https://www.isca-archive.org/interspeech_2018/hara18_interspeech.html"),
        paper("Turn-Taking Predictions Across Languages and Genres Using an LSTM Recurrent Neural Network", "Nigel G. Ward et al.", "SLT", 2018, "https://doi.org/10.1109/SLT.2018.8639633"),
        paper("An Incremental Turn-Taking Model for Task-Oriented Dialog Systems", "Andrei C. Coman et al.", "Interspeech", 2019, "https://www.isca-archive.org/interspeech_2019/coman19_interspeech.html"),
        paper("Turn-Taking Prediction Based on Detection of Transition Relevance Place", "Kohei Hara et al.", "Interspeech", 2019, "https://www.isca-archive.org/interspeech_2019/hara19_interspeech.html"),
        paper("Oh, Jeez! or Uh-Huh? A Listener-Aware Backchannel Predictor on ASR Transcriptions", "Daniel Ortega et al.", "ICASSP", 2020, "https://doi.org/10.1109/ICASSP40776.2020.9054344"),
        paper("TurnGPT: A Transformer-Based Language Model for Predicting Turn-Taking in Spoken Dialog", "Erik Ekstedt and Gabriel Skantze", "EMNLP Findings", 2020, "https://aclanthology.org/2020.findings-emnlp.268/"),
        paper("BPM_MT: Enhanced Backchannel Prediction Model Using Multi-Task Learning", "Jin Yea Jang et al.", "EMNLP", 2021, "https://aclanthology.org/2021.emnlp-main.277/"),
        paper("Projection of Turn Completion in Incremental Spoken Dialogue Systems", "Erik Ekstedt and Gabriel Skantze", "SIGDIAL", 2021, "https://aclanthology.org/2021.sigdial-1.45/"),
        paper("Timing Generating Networks: Neural Network Based Precise Turn-Taking Timing Prediction in Multiparty Conversation", "Shinya Fujie et al.", "Interspeech", 2021, "https://www.isca-archive.org/interspeech_2021/fujie21_interspeech.html"),
        paper("Contextual Acoustic Barge-In Classification for Spoken Dialog Systems", "Dhanush Bekal et al.", "Interspeech", 2022, "https://www.isca-archive.org/interspeech_2022/bekal22_interspeech.html"),
        paper("Gated Multimodal Fusion with Contrastive Learning for Turn-Taking Prediction in Human-Robot Dialogue", "Jiudong Yang et al.", "ICASSP", 2022, "https://doi.org/10.1109/ICASSP43922.2022.9746613"),
        paper("Response Timing Estimation for Spoken Dialog System Using Dialog Act Estimation", "Jin Sakuma et al.", "Interspeech", 2022, "https://www.isca-archive.org/interspeech_2022/sakuma22_interspeech.html"),
        paper("Turn-Taking Prediction for Natural Conversational Speech", "Shuo-Yiin Chang et al.", "Interspeech", 2022, "https://www.isca-archive.org/interspeech_2022/chang22_interspeech.html"),
        paper("Voice Activity Projection: Self-Supervised Learning of Turn-Taking Events", "Erik Ekstedt and Gabriel Skantze", "Interspeech", 2022, "https://www.isca-archive.org/interspeech_2022/ekstedt22_interspeech.html"),
    ]),
    ("Evaluation & Benchmarks", "Evaluation Focus: Full-duplex Interaction", []),
    ("Evaluation & Benchmarks", "Evaluation Focus: Multi-turn Dialogue Capabilities", [
        paper("PARADISE: A Framework for Evaluating Spoken Dialogue Agents", "Marilyn A. Walker et al.", "ACL-EACL", 1997, "https://aclanthology.org/P97-1035/"),
        paper("Quantitative and Qualitative Evaluation of DARPA Communicator Spoken Dialogue Systems", "Julie E. Boland et al.", "ACL", 2001, "https://aclanthology.org/P01-1066/"),
        paper("The PARADISE Evaluation Framework: Issues and Findings", "Melita Hajdinjak and France Mihelic", "Computational Linguistics", 2006, "https://aclanthology.org/J06-2004/"),
        paper("A Framework for Model-Based Evaluation of Spoken Dialog Systems", "Sebastian Moller et al.", "SIGDIAL", 2008, "https://aclanthology.org/W08-0128/"),
        paper("User Simulation as Testing for Spoken Dialog Systems", "Hua Ai and Fuliang Weng", "SIGDIAL", 2008, "https://aclanthology.org/W08-0126/"),
        paper("Spoken Dialog Challenge 2010: Comparison of Live and Control Test Results", "Alan W. Black et al.", "SIGDIAL", 2011, "https://aclanthology.org/W11-2002/"),
        paper("Position Paper: Towards Standardized Metrics and Tools for Spoken and Multimodal Dialog System Evaluation", "Sebastian Moller et al.", "NAACL Workshop", 2012, "https://aclanthology.org/W12-1803/"),
        paper("The Dialog State Tracking Challenge", "Jason D. Williams et al.", "SIGDIAL", 2013, "https://aclanthology.org/W13-4065/"),
        paper("The Second Dialog State Tracking Challenge", "Matthew Henderson et al.", "SIGDIAL", 2014, "https://aclanthology.org/W14-4337/"),
    ]),
    ("Models", "Interaction Mode: Turn-taking", [
        paper("A Robust System for Natural Spoken Dialogue", "James F. Allen et al.", "ACL", 1996, "https://aclanthology.org/P96-1009/"),
        paper("Galaxy-II: A Reference Architecture for Conversational System Development", "Stephanie Seneff et al.", "ICSLP", 1998, "https://www.isca-archive.org/icslp_1998/seneff98b_icslp.html"),
        paper("Let's Go Public! Taking a Spoken Dialog System to the Real World", "Antoine Raux et al.", "Interspeech", 2005, "https://www.isca-archive.org/interspeech_2005/raux05_interspeech.html"),
        paper("The RavenClaw Dialog Management Framework: Architecture and Systems", "Dan Bohus and Alexander I. Rudnicky", "Computer Speech & Language", 2009, "https://doi.org/10.1016/j.csl.2008.10.001"),
        paper("The SEMAINE API: Towards a Standards-Based Framework for Building Emotion-Oriented Systems", "Marc Schroder", "Advances in Human-Computer Interaction", 2010, "https://doi.org/10.1155/2010/319406"),
        paper("Attentive Listening System with Backchanneling, Response Generation and Flexible Turn-Taking", "Divesh Lala et al.", "SIGDIAL", 2017, "https://aclanthology.org/W17-5516/"),
        paper("Towards End-to-End Spoken Dialogue Systems with Turn Embeddings", "Ali Orkan Bayer et al.", "Interspeech", 2017, "https://www.isca-archive.org/interspeech_2017/bayer17_interspeech.html"),
    ]),
    ("Models", "Interaction Mode: Half-duplex / Controlled Barge-in", [
        paper("A Multiparty Multimodal Architecture for Realtime Turntaking", "Kristinn R. Thorisson et al.", "IVA", 2010, "https://doi.org/10.1007/978-3-642-15892-6_37"),
        paper("Towards Incremental Speech Generation in Dialogue Systems", "Gabriel Skantze and Anna Hjalmarsson", "SIGDIAL", 2010, "https://aclanthology.org/W10-4301/"),
        paper("A General, Abstract Model of Incremental Dialogue Processing", "David Schlangen and Gabriel Skantze", "Dialogue & Discourse", 2011, "https://aclanthology.org/2011.dnd-2.11/"),
        paper("A Demonstration of Incremental Speech Understanding and Confidence Estimation in a Virtual Human Dialogue System", "David DeVault and David Traum", "SIGDIAL", 2012, "https://aclanthology.org/W12-1618/"),
        paper("The InproTK 2012 Release", "Timo Baumann and David Schlangen", "NAACL Workshop", 2012, "https://aclanthology.org/W12-1814/"),
        paper("Incremental Dialog Processing in a Task-Oriented Dialog", "Fabrizio Ghigi et al.", "Interspeech", 2014, "https://www.isca-archive.org/interspeech_2014/ghigi14_interspeech.html"),
        paper("An Easy Method to Make Dialogue Systems Incremental", "Hatim Khouzaimi et al.", "SIGDIAL", 2014, "https://aclanthology.org/W14-4314/"),
    ]),
    ("Models", "Interaction Mode: Full-duplex", []),
    ("Training Methods", "Interaction Mode: Turn-taking", [
        paper("NJFun: A Reinforcement Learning Spoken Dialogue System", "Diane Litman et al.", "ANLP-NAACL Workshop", 2000, "https://aclanthology.org/W00-0304/"),
        paper("Optimizing Dialogue Management with Reinforcement Learning: Experiments with the NJFun System", "Satinder Singh et al.", "Journal of Artificial Intelligence Research", 2002, "https://doi.org/10.1613/jair.859"),
        paper("Partially Observable Markov Decision Processes for Spoken Dialog Systems", "Jason D. Williams and Steve Young", "Computer Speech & Language", 2007, "https://doi.org/10.1016/j.csl.2006.06.008"),
        paper("Agenda-Based User Simulation for Bootstrapping a POMDP Dialogue System", "Jost Schatzmann et al.", "HLT-NAACL", 2007, "https://aclanthology.org/N07-2038/"),
        paper("Learning from Real Users: Rating Dialogue Success with Neural Networks for Reinforcement Learning in Spoken Dialogue Systems", "Pei-Hao Su et al.", "Interspeech", 2015, "https://www.isca-archive.org/interspeech_2015/su15_interspeech.html"),
        paper("A Sequence-to-Sequence Model for User Simulation in Spoken Dialogue Systems", "Layla El Asri et al.", "Interspeech", 2016, "https://www.isca-archive.org/interspeech_2016/asri16_interspeech.html"),
        paper("End-to-End LSTM-Based Dialog Control Optimized with Supervised and Reinforcement Learning", "Jason D. Williams and Geoffrey Zweig", "arXiv", 2016, "https://arxiv.org/abs/1606.01269"),
        paper("On-Line Active Reward Learning for Policy Optimisation in Spoken Dialogue Systems", "Pei-Hao Su et al.", "ACL", 2016, "https://arxiv.org/abs/1605.07669"),
        paper("Deep Reinforcement Learning of Dialogue Policies with Less Weight Updates", "Heriberto Cuayahuitl and Seunghak Yu", "Interspeech", 2017, "https://www.isca-archive.org/interspeech_2017/cuayahuitl17_interspeech.html"),
        paper("Towards Generalized Models for Task-Oriented Dialogue Modeling on Spoken Conversations", "Ruijie Yan et al.", "arXiv", 2022, "https://arxiv.org/abs/2203.04045"),
    ]),
    ("Training Methods", "Interaction Mode: Half-duplex / Streaming", [
        paper("Inverse Reinforcement Learning for Micro-Turn Management", "Dongho Kim et al.", "Interspeech", 2014, "https://www.isca-archive.org/interspeech_2014/kim14b_interspeech.html"),
        paper("An Incremental Turn-Taking Model with Active System Barge-In for Spoken Dialog Systems", "Tiancheng Zhao et al.", "SIGDIAL", 2015, "https://aclanthology.org/W15-4606/"),
        paper("Reinforcement Learning for Turn-Taking Management in Incremental Spoken Dialogue Systems", "Hatim Khouzaimi et al.", "IJCAI", 2016, "https://www.ijcai.org/Proceedings/16/Papers/403.pdf"),
        paper("A Methodology for Turn-Taking Capabilities Enhancement in Spoken Dialogue Systems Using Reinforcement Learning", "Hatim Khouzaimi et al.", "Computer Speech & Language", 2018, "https://doi.org/10.1016/j.csl.2017.07.001"),
    ]),
    ("Training Methods", "Interaction Mode: Full-duplex", []),
]


def effective_sections():
    result = [(top, subgroup, []) for top, subgroup, _ in SECTIONS]
    targets = {(top, subgroup): items for top, subgroup, items in result}
    for top, subgroup, items in SECTIONS:
        for item in items:
            destination = RECLASSIFICATIONS.get(item["title"], (top, subgroup))
            if destination not in targets:
                raise ValueError(f"Unknown reclassification target: {destination}")
            targets[destination].append(item)
    return result


def normalize_title(value: str) -> str:
    value = value.casefold().replace("’", "'").replace("–", "-").replace("—", "-")
    return re.sub(r"[^a-z0-9]+", "", value)


def existing_titles(path: Path) -> set[str]:
    if not path.exists():
        return set()
    found = set()
    doc = Document(path)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = re.match(r'^[0-9]+\.?\s*["“](.+?)["”]\.?\s', text)
        if not match:
            match = re.match(r'^["“](.+?)["”]\.?\s', text)
        if match:
            found.add(normalize_title(match.group(1)))
    return found


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.extend([tabs, ind])
    level.extend([start, num_fmt, level_text, suffix, ppr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def number_paragraph(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def add_field(paragraph, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, text, end):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 22, "16324F", 0, 6),
        ("Heading 1", 16, "16324F", 12, 4),
        ("Heading 2", 12.5, "2F5D62", 8, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Paper Entry" not in styles:
        entry = styles.add_style("Paper Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        entry = styles["Paper Entry"]
    entry.base_style = styles["Normal"]
    entry.paragraph_format.space_after = Pt(3)
    entry.paragraph_format.keep_together = True

    if "Scope Note" not in styles:
        note = styles.add_style("Scope Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Scope Note"]
    note.base_style = styles["Normal"]
    note.font.size = Pt(9)
    note.font.italic = True
    note.font.color.rgb = RGBColor(90, 90, 90)
    note.paragraph_format.space_after = Pt(5)


def add_paper_entry(doc: Document, item: dict, num_id: int):
    paragraph = doc.add_paragraph(style="Paper Entry")
    number_paragraph(paragraph, num_id)
    title_run = paragraph.add_run(f'"{item["title"]}". ')
    title_run.bold = True
    author_run = paragraph.add_run(f'{item["authors"]}. ')
    author_run.italic = True
    paragraph.add_run(f'{item["venue"]} {item["year"]}. ')
    add_hyperlink(paragraph, "[Paper]", item["url"])


def build():
    seen = existing_titles(BASE_DOCX) if DEDUPLICATE_AGAINST_BASE else set()
    sections = effective_sections()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)

    title = doc.add_paragraph(style="Title")
    title.add_run(DOCUMENT_TITLE)
    intro = doc.add_paragraph(INTRO_TEXT)
    intro.paragraph_format.space_after = Pt(4)
    note = doc.add_paragraph(style="Scope Note")
    note.add_run("Scope: multi-turn spoken interaction, spoken dialogue corpora and evaluation, turn-taking/interruption, incremental or barge-in-capable systems, and training methods for spoken dialogue. Text-only dialogue papers and single-turn ASR/TTS/SLU work are excluded unless they directly model turn-taking for spoken systems.")

    # Compact summary band.
    total_candidates = sum(sum(item["year"] >= MIN_YEAR for item in items) for _, _, items in sections)
    kept_candidates = sum(
        1
        for _, _, items in sections
        for item in items
        if item["year"] >= MIN_YEAR and normalize_title(item["title"]) not in seen
    )
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(5.6), Cm(5.6), Cm(5.6)]
    third_text = SUMMARY_THIRD_TEXT or f"{total_candidates - kept_candidates} existing title(s) removed"
    labels = [("Cutoff", CUTOFF_TEXT), (COUNT_LABEL, str(kept_candidates)), (SUMMARY_THIRD_LABEL, third_text)]
    for index, (label, value) in enumerate(labels):
        cell = table.cell(0, index)
        cell.width = widths[index]
        set_cell_margins(cell)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "EAF1F4")
        cell._tc.get_or_add_tcPr().append(shade)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "\n")
        r1.bold = True
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor(47, 93, 98)
        r2 = p.add_run(value)
        r2.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    last_top = None
    for top, subgroup, items in sections:
        filtered = [x for x in items if x["year"] >= MIN_YEAR and normalize_title(x["title"]) not in seen]
        filtered.sort(key=lambda x: (x["year"], x["title"].casefold()))
        if top != last_top:
            doc.add_heading(top, level=1)
            last_top = top
        if subgroup:
            doc.add_heading(subgroup, level=2)
        if not filtered:
            note = doc.add_paragraph(style="Scope Note")
            note.add_run(EMPTY_SECTION_TEXT)
            continue
        num_id = add_numbering_definition(doc)
        for item in filtered:
            add_paper_entry(doc, item, num_id)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = HEADER_TEXT
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(110, 110, 110)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer, "PAGE")
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(110, 110, 110)

    properties = doc.core_properties
    properties.title = DOCUMENT_TITLE
    properties.subject = SUBJECT_TEXT
    properties.author = "Codex"
    properties.keywords = "spoken dialogue, turn-taking, barge-in, incremental dialogue, dataset, benchmark"
    doc.save(OUTPUT_DOCX)
    print(f"saved={OUTPUT_DOCX}")
    print(f"existing_titles={len(seen)}")
    print(f"candidate_papers={total_candidates}")
    print(f"written_papers={kept_candidates}")


if __name__ == "__main__":
    build()
