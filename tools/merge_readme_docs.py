from collections import OrderedDict
from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"C:\Users\26547\Desktop\read\readme.docx")
ADD = Path(r"C:\Users\26547\Desktop\read\readme_add2.docx")
SUPPLEMENT = Path(r"D:\vscode\MT-Speech-Dialogue\readme_2018_to_2022.docx")
OUTPUT = BASE
README = Path(r"D:\vscode\MT-Speech-Dialogue\README.md")

SECTION_ORDER = [
    "Survey",
    "General Surveys",
    "Datasets",
    "Evaluation & Benchmarks",
    "Models",
    "Training Methods",
]

SUBGROUP_ORDER = {
    "Evaluation & Benchmarks": [
        "Evaluation Focus: Turn-taking & Interruption",
        "Evaluation Focus: Full-duplex Interaction",
        "Evaluation Focus: Multi-turn Dialogue Capabilities",
    ],
    "Models": [
        "Interaction Mode: Turn-taking",
        "Interaction Mode: Half-duplex / Controlled Barge-in",
        "Interaction Mode: Full-duplex",
    ],
    "Training Methods": [
        "Interaction Mode: Turn-taking",
        "Interaction Mode: Half-duplex / Streaming",
        "Interaction Mode: Full-duplex",
    ],
}

BASE_NUM_IDS = {
    "Survey": 7,
    "General Surveys": 8,
    "Datasets": 9,
    "Evaluation & Benchmarks": 10,
    "Models": 11,
    "Training Methods": 12,
}

# Entries from readme_add2.docx that survive topic, category, link, and
# duplicate review. Titles are normalized to the official publication title.
APPROVED_ADDITIONS = [
    {
        "section": "Survey",
        "subgroup": "",
        "title": "Turn-Taking Modelling in Conversational Systems: A Review of Recent Advances",
        "author": "Rutherford Agbeshi Patamia et al.",
        "venue": "Technologies 2025",
        "year": 2025,
        "url": "https://doi.org/10.3390/technologies13120591",
    },
    {
        "section": "Models",
        "subgroup": "Interaction Mode: Full-duplex",
        "title": "Duplex Conversation in Outbound Agent System",
        "author": "Chunxiang Jin et al.",
        "venue": "Interspeech 2021",
        "year": 2021,
        "url": "https://www.isca-archive.org/interspeech_2021/jin21b_interspeech.html",
    },
    {
        "section": "Models",
        "subgroup": "Interaction Mode: Full-duplex",
        "title": "Language Model Can Listen While Speaking",
        "author": "Ziyang Ma et al.",
        "venue": "AAAI 2025",
        "year": 2025,
        "url": "https://arxiv.org/abs/2408.02622",
    },
    {
        "section": "Datasets",
        "subgroup": "",
        "title": "DialogueAgents: A Hybrid Agent-Based Speech Synthesis Framework for Multi-Party Dialogue",
        "author": "Zheyuan Zhang et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "url": "https://arxiv.org/abs/2504.14482",
    },
    {
        "section": "Datasets",
        "subgroup": "",
        "title": "Open-Source Full-Duplex Conversational Datasets for Natural and Interactive Speech Synthesis",
        "author": "Rui Liu et al.",
        "venue": "arXiv 2025",
        "year": 2025,
        "url": "https://arxiv.org/abs/2509.04093",
    },
    {
        "section": "Models",
        "subgroup": "Interaction Mode: Turn-taking",
        "title": "Voice Activity Projection: Self-supervised Learning of Turn-taking Events",
        "author": "Erik Ekstedt and Gabriel Skantze",
        "venue": "Interspeech 2022",
        "year": 2022,
        "url": "https://doi.org/10.21437/Interspeech.2022-10955",
    },
    {
        "section": "Models",
        "subgroup": "Interaction Mode: Half-duplex / Controlled Barge-in",
        "title": "VITA: Towards Open-Source Interactive Omni Multimodal LLM",
        "author": "Chaoyou Fu et al.",
        "venue": "arXiv 2024",
        "year": 2024,
        "url": "https://arxiv.org/abs/2408.05211",
    },
]

# Reviewed moves based on each paper's primary contribution. Interaction-mode
# labels are only applied to Models and Training Methods, not to datasets.
RECLASSIFICATIONS = {
    "A Review of Dialogue Systems: From Trained Monkeys to Stochastic Parrots": ("Survey", ""),
    "On The Landscape of Spoken Language Models: A Comprehensive Survey": ("General Surveys", ""),
    "Recent Advances in Speech Language Models: A Survey": ("General Surveys", ""),
    "Let's Go Real Talk: Spoken Dialogue Model for Face-to-Face Conversation": ("Models", "Interaction Mode: Turn-taking"),
    "MMedFD: A Real-world Healthcare Benchmark for Multi-turn Full-Duplex Automatic Speech Recognition": ("Evaluation & Benchmarks", "Evaluation Focus: Full-duplex Interaction"),
    "Data-Centric Improvements for Enhancing Multi-Modal Understanding in Spoken Conversation Modeling (ASK-QA)": ("Training Methods", "Interaction Mode: Turn-taking"),
    "Using Transition Duration to Improve Turn-taking in Conversational Agents": ("Models", "Interaction Mode: Turn-taking"),
    "When can I Speak? Predicting Initiation Points for Spoken Dialogue Agents": ("Models", "Interaction Mode: Turn-taking"),
    "Device Directedness with Contextual Cues for Spoken Dialog Systems": ("Models", "Interaction Mode: Half-duplex / Controlled Barge-in"),
    "Triadic Multi-party Voice Activity Projection for Turn-taking in Spoken Dialogue Systems": ("Models", "Interaction Mode: Turn-taking"),
    "Mini-Omni: Language Models Can Hear, Talk While Thinking in Streaming": ("Models", "Interaction Mode: Turn-taking"),
    "GLM-4-Voice: Towards Intelligent and Human-Like End-to-End Spoken Chatbot": ("Models", "Interaction Mode: Turn-taking"),
    "Qwen2.5-Omni Technical Report": ("Models", "Interaction Mode: Turn-taking"),
    "LLaMA-Omni 2: LLM-based Real-time Spoken Chatbot with Autoregressive Streaming Speech Synthesis": ("Models", "Interaction Mode: Turn-taking"),
    "FLM-Audio: Natural Monologues Improves Native Full-Duplex Chatbots via Dual Training": ("Training Methods", "Interaction Mode: Full-duplex"),
    "SHANKS: Simultaneous Hearing and Thinking for Spoken Language Models": ("Models", "Interaction Mode: Half-duplex / Controlled Barge-in"),
    "Chronological Thinking in Full-Duplex Spoken Dialogue Language Models": ("Models", "Interaction Mode: Full-duplex"),
    "LLM-Enhanced Dialogue Management for Full-Duplex Spoken Dialogue Systems": ("Models", "Interaction Mode: Full-duplex"),
}

SURVEY_COMPARISON_ROWS = [
    # General Surveys first, followed by core surveys in chronological order.
    ("Serban et al. (2018)", "General", "△", "-", "-", "✓", "✓", "△", "△", "✓", "△"),
    ("Gao et al. (2019)", "General", "△", "-", "-", "△", "△", "✓", "✓", "✓", "-"),
    ("Balaraman et al. (2021)", "General", "△", "-", "-", "✓", "✓", "✓", "△", "✓", "✓"),
    ("Gu et al. (2022)", "General", "△", "△", "-", "△", "△", "✓", "△", "✓", "-"),
    ("Latif et al. (2023)", "General", "△", "-", "-", "△", "△", "✓", "△", "✓", "-"),
    ("Winata et al. (2024)", "General", "-", "-", "-", "✓", "✓", "△", "✓", "✓", "△"),
    ("Peng et al. (2024)", "General", "△", "-", "-", "△", "✓", "✓", "✓", "✓", "△"),
    ("Arora et al. (2025)", "General", "△", "△", "△", "△", "✓", "✓", "✓", "✓", "△"),
    ("Cui et al. (2025)", "General", "△", "-", "△", "△", "✓", "✓", "✓", "✓", "△"),
    ("Yu et al. (2025)", "General", "-", "-", "-", "✓", "✓", "△", "✓", "✓", "✓"),
    ("Skantze (2021)", "Core", "✓", "✓", "△", "△", "△", "✓", "-", "✓", "-"),
    ("Patlan et al. (2023)", "Core", "✓", "△", "-", "△", "✓", "✓", "✓", "✓", "△"),
    ("Ji et al. / WavChat (2024)", "Core", "✓", "✓", "✓", "✓", "✓", "✓", "✓", "✓", "✓"),
    ("Chen and Yu (2025)", "Core", "✓", "✓", "✓", "△", "✓", "✓", "△", "✓", "✓"),
    ("Castillo-López et al. (2025)", "Core", "✓", "✓", "-", "✓", "△", "✓", "△", "✓", "△"),
    ("Patamia et al. (2025)", "Core", "✓", "✓", "△", "✓", "✓", "✓", "△", "✓", "△"),
    ("Lu et al. (2026)", "Core", "✓", "✓", "✓", "✓", "✓", "✓", "△", "✓", "✓"),
    ("Current Collection", "Current", "✓", "✓", "✓", "✓", "✓", "✓", "✓", "✓", "✓"),
]

METADATA_CORRECTIONS = {
    "Turn-taking in Conversational Systems and Human-Robot Interaction: A Review": {
        "author": "Gabriel Skantze",
    },
    "From Turn-Taking to Synchronous Dialogue: A Survey of Full-Duplex Spoken Language Models": {
        "author": "Yuxuan Chen and Haoyuan Yu",
        "venue": "arXiv 2025",
        "year": 2025,
    },
    "A Survey of Recent Advances on Turn-taking Modeling in Spoken Dialogue Systems": {
        "author": "Galo Castillo-López et al.",
    },
    "A Survey of Full-Duplex Spoken Dialogue Systems: Architectural Hierarchy, Interaction Ontology, and Decision State Machine": {
        "author": "Jingyu Lu et al.",
    },
    "A Survey of Available Corpora For Building Data-Driven Dialogue Systems: The Journal Version": {
        "author": "Iulian V. Serban et al.",
    },
    "Neural Approaches to Conversational AI": {
        "author": "Jianfeng Gao et al.",
    },
    "Recent Neural Methods on Dialogue State Tracking for Task-Oriented Dialogue Systems: A Survey": {
        "author": "Vevake Balaraman et al.",
    },
    "Who Says What to Whom: A Survey of Multi-Party Conversations": {
        "author": "Jia-Chen Gu et al.",
    },
    "A Survey on Speech Large Language Models for Understanding": {
        "title": "A Survey on Speech Large Language Models",
        "author": "Jing Peng et al.",
    },
}


def normalize_title(title):
    return re.sub(r"[^a-z0-9]+", "", title.lower())


def title_from_text(text):
    if text.startswith('"') and '". ' in text:
        return text.split('". ', 1)[0][1:]
    return None


def hyperlink_from_paragraph(paragraph):
    for node in paragraph._p.findall(".//" + qn("w:instrText")):
        match = re.search(r'HYPERLINK\s+"([^"]+)"', node.text or "")
        if match:
            return match.group(1)
    for link in paragraph._p.findall(".//" + qn("w:hyperlink")):
        rel_id = link.get(qn("r:id"))
        if rel_id and rel_id in paragraph.part.rels:
            return paragraph.part.rels[rel_id].target_ref
    return ""


def parse_entry(paragraph, source_rank, source_index):
    text = paragraph.text.strip()
    title = title_from_text(text)
    if not title:
        return None
    tail = text.split('". ', 1)[1]
    # Source documents use both "Paper" and "[Paper]" hyperlink labels. Strip
    # every trailing display label before parsing author and venue metadata.
    tail = re.sub(r"(?:\s*\.?\s*\[?Paper\]?)+\s*\.?\s*$", "", tail, flags=re.IGNORECASE).rstrip()
    year_match = re.search(r"\b(20\d{2})\b", tail)
    year = int(year_match.group(1)) if year_match else 9999
    italic_text = "".join(run.text for run in paragraph.runs if run.italic).strip()
    author = italic_text or tail.split(". ", 1)[0].strip()
    after_author = tail[len(author):].lstrip(". ") if tail.startswith(author) else tail
    venue = after_author.rstrip(". ").strip()
    return {
        "title": title,
        "author": author,
        "venue": venue,
        "year": year,
        "url": hyperlink_from_paragraph(paragraph),
        "source_rank": source_rank,
        "source_index": source_index,
    }


def extract(document, source_rank):
    result = OrderedDict()
    section = ""
    subgroup = ""
    index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 1":
            section = text
            subgroup = ""
            result.setdefault(section, OrderedDict())
        elif paragraph.style.name == "Heading 2":
            subgroup = text
            result.setdefault(section, OrderedDict()).setdefault(subgroup, [])
        else:
            entry = parse_entry(paragraph, source_rank, index)
            if entry:
                result.setdefault(section, OrderedDict()).setdefault(subgroup, []).append(entry)
                index += 1
    return result


def apply_reviewed_corrections(data):
    corrected = OrderedDict()
    for section, groups in data.items():
        for subgroup, entries in groups.items():
            for entry in entries:
                entry = dict(entry)
                entry.update(METADATA_CORRECTIONS.get(entry["title"], {}))
                target_section, target_subgroup = RECLASSIFICATIONS.get(
                    entry["title"], (section, subgroup)
                )
                corrected.setdefault(target_section, OrderedDict()).setdefault(
                    target_subgroup, []
                ).append(entry)
    return corrected


def merge(*sources):
    merged = OrderedDict()
    seen = set()
    duplicate_titles = []
    for section in SECTION_ORDER:
        merged[section] = OrderedDict()
        subgroup_names = SUBGROUP_ORDER.get(section, [""])
        for subgroup in subgroup_names:
            entries = []
            for data in sources:
                entries.extend(data.get(section, {}).get(subgroup, []))
            unique = []
            for entry in entries:
                key = normalize_title(entry["title"])
                if key in seen:
                    duplicate_titles.append(entry["title"])
                    continue
                seen.add(key)
                unique.append(entry)
            unique.sort(key=lambda item: (item["year"], item["source_rank"], item["source_index"]))
            merged[section][subgroup] = unique
    return merged, duplicate_titles


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def next_number_id(document, source_id):
    numbering = document.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    source = next(node for node in nums if node.get(qn("w:numId")) == str(source_id))
    new_id = max(int(node.get(qn("w:numId"))) for node in nums) + 1
    clone = deepcopy(source)
    clone.set(qn("w:numId"), str(new_id))
    for override in clone.findall(qn("w:lvlOverride")):
        clone.remove(override)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    clone.append(override)
    numbering.append(clone)
    return str(new_id)


def set_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), num_id)
    num_pr.extend([ilvl, num])
    ppr.insert(0, num_pr)


def add_hyperlink_field(paragraph, url, label="Paper"):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' HYPERLINK "{url}" \\h '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    display_text = OxmlElement("w:t")
    display_text.text = label
    display_run.extend([rpr, display_text])
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, end):
        wrapper = OxmlElement("w:r")
        wrapper.append(node)
        paragraph._p.append(wrapper)
        if node is separate:
            paragraph._p.append(display_run)


def add_paper(document, entry, num_id):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.widow_control = True
    set_numbering(paragraph, num_id)
    paragraph.add_run(f'"{entry["title"]}". ')
    author = paragraph.add_run(f'{entry["author"]} ')
    author.italic = True
    paragraph.add_run(f'{entry["venue"]}. ')
    add_hyperlink_field(paragraph, entry["url"])


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, value=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def format_table_text(cell, bold=False, color="1F1F1F", size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(color)


def set_table_geometry(table, widths):
    table.autofit = False
    table_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)


def append_survey_comparison(document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    heading = document.add_paragraph("Survey Coverage Comparison", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(3)

    note = document.add_paragraph(style="Normal")
    note.paragraph_format.space_after = Pt(3)
    label = note.add_run("Classification note. ")
    label.bold = True
    note.add_run(
        "Patlan et al. (2023) is classified as a core survey because it explicitly reviews "
        "spoken dialogue systems and multi-turn human-agent interaction. General surveys "
        "cover adjacent areas with broader scopes."
    )
    for run in note.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)

    caption = document.add_paragraph(style="Normal")
    caption.paragraph_format.space_after = Pt(2)
    run = caption.add_run(
        "Table 1. Coverage comparison of surveys related to Multi-Turn Speech Dialogue."
    )
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)

    legend = document.add_paragraph(style="Normal")
    legend.paragraph_format.space_after = Pt(4)
    legend.add_run(
        "General Surveys are listed first, followed by core surveys. "
        "✓ = systematic coverage; △ = partial or supporting coverage; - = not a substantive focus. "
        "Current Collection refers to this curated review."
    )
    for run in legend.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

    table = document.add_table(rows=2, cols=10)
    table.style = "Table Grid"
    widths = [3000, 1373, 1373, 1373, 1373, 1373, 1373, 1373, 1373, 1376]
    set_table_geometry(table, widths)

    top = table.rows[0].cells
    top[0].text = "Survey"
    top[0].merge(table.rows[1].cells[0])
    top[1].text = "Research Scope"
    top[1].merge(top[3])
    top[4].text = "Content Coverage"
    top[4].merge(top[7])
    top[8].text = "Organization"
    top[8].merge(top[9])

    labels = [
        "Multi-turn\nSpeech",
        "Turn-taking &\nInterruption",
        "Full-duplex",
        "Datasets",
        "Evaluation",
        "Models",
        "Training\nMethods",
        "Interaction\nTaxonomy",
        "Benchmark",
    ]
    for index, text in enumerate(labels, 1):
        table.rows[1].cells[index].text = text

    for row in table.rows[:2]:
        set_repeat_table_header(row)
        for cell in row.cells:
            set_cell_shading(cell, "D9E5F3")
            format_table_text(cell, bold=True, color="17365D", size=7.2)

    for row_index, values in enumerate(SURVEY_COMPARISON_ROWS, start=2):
        row = table.add_row()
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
        citation, group, *marks = values
        row.cells[0].text = citation
        for index, mark in enumerate(marks, 1):
            row.cells[index].text = mark
        fill = "F2F2F2" if row_index % 2 == 0 else "FFFFFF"
        if group == "Current":
            fill = "D9EAD3"
        elif group == "Core":
            fill = "EEF4FB" if row_index % 2 == 0 else "FFFFFF"
        for index, cell in enumerate(row.cells):
            set_cell_shading(cell, fill)
            format_table_text(
                cell,
                bold=group == "Current",
                size=7.1,
                align=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )


def build_docx(merged):
    document = Document(BASE)
    clear_body(document)
    document.add_paragraph("Multi-Turn Speech Dialogue", style="Title")
    document.add_paragraph("A curated paper list for Multi-Turn Speech Dialogue.", style="Normal")
    for section, groups in merged.items():
        heading = document.add_paragraph(section, style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        for subgroup, entries in groups.items():
            if subgroup:
                subheading = document.add_paragraph(subgroup, style="Heading 2")
                subheading.paragraph_format.keep_with_next = True
            num_id = next_number_id(document, BASE_NUM_IDS[section])
            for entry in entries:
                add_paper(document, entry, num_id)
    append_survey_comparison(document)
    document.save(OUTPUT)


def escape_md(text):
    return text.replace("|", "\\|")


def build_markdown(merged):
    lines = ["# MT-Speech-Dialogue", "", "> A curated paper list for **Multi-Turn Speech Dialogue**.", ""]
    for section, groups in merged.items():
        lines.extend([f"## {section}", ""])
        for subgroup, entries in groups.items():
            if subgroup:
                lines.extend([f"### {subgroup}", ""])
            for index, entry in enumerate(entries, 1):
                lines.append(
                    f'{index}. **"{escape_md(entry["title"])}"**. '
                    f'*{escape_md(entry["author"])}* {escape_md(entry["venue"])}. '
                    f'[[Paper]({entry["url"]})]'
                )
            lines.append("")
    README.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def main():
    base_doc = Document(BASE)
    if not SUPPLEMENT.exists():
        raise FileNotFoundError(SUPPLEMENT)
    reviewed = OrderedDict()
    for index, item in enumerate(APPROVED_ADDITIONS):
        entry = dict(item)
        section = entry.pop("section")
        subgroup = entry.pop("subgroup")
        entry.update(source_rank=1, source_index=index)
        reviewed.setdefault(section, OrderedDict()).setdefault(subgroup, []).append(entry)
    base_data = apply_reviewed_corrections(extract(base_doc, 0))
    supplement_data = apply_reviewed_corrections(extract(Document(SUPPLEMENT), 2))
    merged, duplicates = merge(base_data, reviewed, supplement_data)
    build_docx(merged)
    build_markdown(merged)
    total = sum(len(entries) for groups in merged.values() for entries in groups.values())
    print(f"Merged {total} papers into {OUTPUT}")
    print(f"Updated {README}")
    print(f"Skipped duplicate titles: {len(duplicates)}")
    for title in duplicates:
        print(f"  - {title}")


if __name__ == "__main__":
    main()
