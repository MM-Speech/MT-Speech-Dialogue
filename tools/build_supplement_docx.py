from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\Users\26547\Desktop\read\readme_restructured.docx")
OUTPUT = Path(r"D:\vscode\MT-Speech-Dialogue\readme_supplement_2022_present.docx")


def paper(date, title, author, venue, url):
    return {
        "date": date,
        "title": title,
        "author": author,
        "venue": venue,
        "url": url,
    }


SECTIONS = [
    (
        "Survey",
        None,
        [
            paper(
                "2025-05",
                "A Survey of Recent Advances on Turn-taking Modeling in Spoken Dialogue Systems",
                "Galo Castillo-López et al.",
                "IWSDS 2025",
                "https://aclanthology.org/2025.iwsds-1.27/",
            ),
        ],
    ),
    (
        "Datasets",
        None,
        [
            paper(
                "2022-07-03",
                "DailyTalk: Spoken Dialogue Dataset for Conversational Text-to-Speech",
                "Keon Lee et al.",
                "arXiv 2022",
                "https://arxiv.org/abs/2207.01063",
            ),
            paper(
                "2022-12",
                "DOROTHIE: Spoken Dialogue for Handling Unexpected Situations in Interactive Autonomous Driving Agents",
                "Ziqiao Ma et al.",
                "EMNLP Findings 2022",
                "https://aclanthology.org/2022.findings-emnlp.354/",
            ),
            paper(
                "2023-05-22",
                "SpokenWOZ: A Large-Scale Speech-Text Benchmark for Spoken Task-Oriented Dialogue Agents",
                "Shuzheng Si et al.",
                "arXiv 2023",
                "https://arxiv.org/abs/2305.13040",
            ),
            paper(
                "2023-08",
                "CALLS: Japanese Empathetic Dialogue Speech Corpus of Complaint Handling and Attentive Listening in Customer Center",
                "Yuki Saito et al.",
                "Interspeech 2023",
                "https://www.isca-archive.org/interspeech_2023/saito23b_interspeech.html",
            ),
            paper(
                "2024-07-22",
                "J-CHAT: Japanese Large-scale Spoken Dialogue Corpus for Spoken Dialogue Language Modeling",
                "Wataru Nakata et al.",
                "arXiv 2024",
                "https://arxiv.org/abs/2407.15828",
            ),
            paper(
                "2024-08",
                "Let's Go Real Talk: Spoken Dialogue Model for Face-to-Face Conversation",
                "Se Park et al.",
                "ACL 2024",
                "https://aclanthology.org/2024.acl-long.860/",
            ),
            paper(
                "2025-04",
                "Behavior-SD: Behaviorally Aware Spoken Dialogue Generation with Large Language Models",
                "Sehun Lee et al.",
                "NAACL 2025",
                "https://aclanthology.org/2025.naacl-long.484/",
            ),
            paper(
                "2025-08-06",
                "RealTalk-CN: A Realistic Chinese Speech-Text Dialogue Benchmark With Cross-Modal Interaction Analysis",
                "Enzhi Wang et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2508.10015",
            ),
            paper(
                "2026-05",
                "Exploring Emotional Nuances in Spoken Dialogue: Dataset Construction and Prediction of Emotional Dialogue Breakdown",
                "Hyuga Nakaguro et al.",
                "IWSDS 2026",
                "https://aclanthology.org/2026.iwsds-1.9/",
            ),
            paper(
                "2026-07-06",
                "DuplexChat: Constructing Speaker-Separated Full-Duplex Dialogue Speech at Scale for Spoken Dialogue Language Modeling",
                "Wataru Nakata et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2607.04941",
            ),
            paper(
                "2026-07",
                "Dial HEALTHDIAL for Advice: A Multilingual and Multi-Parallel Spoken Dialogue Dataset for Knowledge-Grounded Information Seeking",
                "Songbo Hu et al.",
                "ACL Findings 2026",
                "https://aclanthology.org/2026.findings-acl.1275/",
            ),
        ],
    ),
    (
        "Evaluation & Benchmarks",
        "Evaluation Focus: Turn-taking & Interruption",
        [
            paper(
                "2022-09",
                "Using Transition Duration to Improve Turn-taking in Conversational Agents",
                "Charles Threlkeld et al.",
                "SIGDIAL 2022",
                "https://aclanthology.org/2022.sigdial-1.20/",
            ),
            paper(
                "2022-09-01",
                "When can I Speak? Predicting Initiation Points for Spoken Dialogue Agents",
                "Siyan Li et al.",
                "SIGDIAL 2022",
                "https://aclanthology.org/2022.sigdial-1.22/",
            ),
            paper(
                "2022-09-02",
                "How Much Does Prosody Help Turn-taking? Investigations using Voice Activity Projection Models",
                "Erik Ekstedt et al.",
                "SIGDIAL 2022",
                "https://aclanthology.org/2022.sigdial-1.51/",
            ),
            paper(
                "2022-11-23",
                "Device Directedness with Contextual Cues for Spoken Dialog Systems",
                "Dhanush Bekal et al.",
                "arXiv 2022",
                "https://arxiv.org/abs/2211.13280",
            ),
            paper(
                "2025-08",
                "Triadic Multi-party Voice Activity Projection for Turn-taking in Spoken Dialogue Systems",
                "Mikey Elmers et al.",
                "Interspeech 2025",
                "https://www.isca-archive.org/interspeech_2025/elmers25_interspeech.html",
            ),
            paper(
                "2026-05-19",
                "Synchronization and Turn-Taking in Full-Duplex Speech Dialogue Models",
                "Pablo Riera et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2605.20356",
            ),
            paper(
                "2026-07",
                "Still Between Us? Evaluating and Improving Voice Assistant Robustness to Third-Party Interruptions",
                "Dongwook Lee et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-long.1902/",
            ),
            paper(
                "2026-08",
                "Rethinking Binary Evaluation of Turn-Taking under Inherent Ambiguity",
                "Yunosuke Kubo et al.",
                "SIGDIAL 2026",
                "https://aclanthology.org/2026.sigdial-1.2/",
            ),
        ],
    ),
    (
        "Evaluation & Benchmarks",
        "Evaluation Focus: Full-duplex Interaction",
        [
            paper(
                "2025-08",
                "FD-Bench: A Full-Duplex Benchmarking Pipeline Designed for Full Duplex Spoken Dialogue Systems",
                "Yizhou Peng et al.",
                "Interspeech 2025",
                "https://www.isca-archive.org/interspeech_2025/peng25b_interspeech.html",
            ),
            paper(
                "2026-03-14",
                "τ-Voice: Benchmarking Full-Duplex Voice Agents on Real-World Domains",
                "Soham Ray et al.",
                "ICML 2026",
                "https://arxiv.org/abs/2603.13686",
            ),
            paper(
                "2026-04-06",
                "Full-Duplex-Bench-v3: Benchmarking Tool Use for Full-Duplex Voice Agents Under Real-World Disfluency",
                "Guan-Ting Lin et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2604.04847",
            ),
            paper(
                "2026-04-23",
                "Full-Duplex Interaction in Spoken Dialogue Systems: A Comprehensive Study from the ICASSP 2026 HumDial Challenge",
                "Chengyou Wang et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2604.21406",
            ),
            paper(
                "2026-07-01",
                "Full-Duplex-Bench-v2: A Multi-Turn Evaluation Framework for Duplex Dialogue Systems with an Automated Examiner",
                "Guan-Ting Lin et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-short.4/",
            ),
            paper(
                "2026-07-02",
                "MTR-DuplexBench: Towards a Comprehensive Evaluation of Multi-Round Conversations for Full-Duplex Speech Language Models",
                "Zhang He et al.",
                "ACL Findings 2026",
                "https://aclanthology.org/2026.findings-acl.263/",
            ),
            paper(
                "2026-07-31",
                "M3-DuplexBench: A Multi-Turn, Multilingual, Multidomain Benchmark for Full-Duplex Spoken Dialogue Models",
                "Ryo Fukuda et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2607.29125",
            ),
        ],
    ),
    (
        "Evaluation & Benchmarks",
        "Evaluation Focus: Other Dialogue Capabilities (其他对话能力)",
        [
            paper(
                "2024-12-06",
                "Benchmarking Open-ended Audio Dialogue Understanding for Large Audio-Language Models",
                "Kuofeng Gao et al.",
                "arXiv 2024",
                "https://arxiv.org/abs/2412.05167",
            ),
            paper(
                "2025-08-01",
                "SOVA-Bench: Benchmarking the Speech Conversation Ability for LLM-based Voice Assistant",
                "Yixuan Hou et al.",
                "Interspeech 2025",
                "https://www.isca-archive.org/interspeech_2025/hou25b_interspeech.html",
            ),
            paper(
                "2025-08-22",
                "MTalk-Bench: Evaluating Speech-to-Speech Models in Multi-Turn Dialogues via Arena-style and Rubrics Protocols",
                "Yuhao Du et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2508.18240",
            ),
            paper(
                "2025-10-09",
                "VoiceAgentBench: Are Voice Assistants Ready for Agentic Tasks?",
                "Dhruv Jain et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2510.07978",
            ),
            paper(
                "2025-11",
                "C3: A Bilingual Benchmark for Spoken Dialogue Models Exploring Challenges in Complex Conversations",
                "Chengqian Ma et al.",
                "EMNLP 2025",
                "https://aclanthology.org/2025.emnlp-main.1160/",
            ),
            paper(
                "2026-05",
                "The Context Trap: Why End-to-End Audio Language Models Fail Multi-turn Dialogues",
                "Zhi Rui Tam et al.",
                "IWSDS 2026",
                "https://aclanthology.org/2026.iwsds-1.7/",
            ),
            paper(
                "2026-07-01",
                "SDiaReward: Modeling and Benchmarking Spoken Dialogue Rewards with Modality and Colloquialness",
                "Jingyu Lu et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-long.185/",
            ),
            paper(
                "2026-07-02",
                "SpeakerSleuth: Can Large Audio-Language Models Judge Speaker Consistency across Multi-turn Dialogues?",
                "Jonggeun Lee et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-long.944/",
            ),
            paper(
                "2026-07-03",
                "Style Amnesia: Investigating Speaking Style Degradation and Mitigation in Multi-Turn Spoken Language Models",
                "Yu-Xiang Lin et al.",
                "ACL Findings 2026",
                "https://aclanthology.org/2026.findings-acl.304/",
            ),
        ],
    ),
    (
        "Models",
        "Interaction Mode: Turn-taking",
        [
            paper(
                "2023-03",
                "Generative Spoken Dialogue Language Modeling",
                "Tu Anh Nguyen et al.",
                "TACL 2023",
                "https://aclanthology.org/2023.tacl-1.15/",
            ),
            paper(
                "2024-09-10",
                "LLaMA-Omni: Seamless Speech Interaction with Large Language Models",
                "Qingkai Fang et al.",
                "arXiv 2024",
                "https://arxiv.org/abs/2409.06666",
            ),
            paper(
                "2025-02-17",
                "Step-Audio: Unified Understanding and Generation in Intelligent Speech Interaction",
                "Ailin Huang et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2502.11946",
            ),
            paper(
                "2025-02-24",
                "Baichuan-Audio: A Unified Framework for End-to-End Speech Interaction",
                "Tianpeng Li et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2502.17239",
            ),
            paper(
                "2025-04-25",
                "Kimi-Audio Technical Report",
                "KimiTeam et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2504.18425",
            ),
            paper(
                "2025-09-22",
                "Qwen3-Omni Technical Report",
                "Jin Xu et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2509.17765",
            ),
            paper(
                "2025-11-01",
                "PACHAT: Persona-Aware Speech Assistant for Multi-party Dialogue",
                "Dongjie Fu et al.",
                "EMNLP 2025",
                "https://aclanthology.org/2025.emnlp-main.1492/",
            ),
            paper(
                "2026-07-01",
                "VoxMind: An End-to-End Agentic Spoken Dialogue System",
                "Tianle Liang et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-long.459/",
            ),
            paper(
                "2026-07-02",
                "ZipVoice-Dialog: Non-Autoregressive Spoken Dialogue Generation with Flow Matching",
                "Han Zhu et al.",
                "ACL Findings 2026",
                "https://aclanthology.org/2026.findings-acl.1928/",
            ),
        ],
    ),
    (
        "Models",
        "Interaction Mode: Half-duplex / Barge-in",
        [
            paper(
                "2024-10-15",
                "Mini-Omni2: Towards Open-source GPT-4o with Vision, Speech and Duplex Capabilities",
                "Zhifei Xie et al.",
                "arXiv 2024",
                "https://arxiv.org/abs/2410.11190",
            ),
            paper(
                "2025-05-06",
                "VITA-Audio: Fast Interleaved Cross-Modal Token Generation for Efficient Large Speech-Language Model",
                "Zuwei Long et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2505.03739",
            ),
            paper(
                "2025-11-14",
                "AV-Dialog: Spoken Dialogue Models with Audio-Visual Input",
                "Tuochao Chen et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2511.11124",
            ),
            paper(
                "2026-03-16",
                "SoulX-Duplug: Plug-and-Play Streaming State Prediction Module for Realtime Full-Duplex Speech Conversation",
                "Ruiqi Yan et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2603.14877",
            ),
        ],
    ),
    (
        "Models",
        "Interaction Mode: Full-duplex",
        [
            paper(
                "2022-05-30",
                "Duplex Conversation: Towards Human-like Interaction in Spoken Dialogue Systems",
                "Ting-En Lin et al.",
                "arXiv 2022",
                "https://arxiv.org/abs/2205.15060",
            ),
            paper(
                "2024-05-29",
                "A Full-duplex Speech Dialogue Scheme Based On Large Language Models",
                "Peng Wang et al.",
                "NeurIPS 2024",
                "https://arxiv.org/abs/2405.19487",
            ),
            paper(
                "2025-08-01",
                "Towards a Japanese Full-duplex Spoken Dialogue System",
                "Atsumoto Ohashi et al.",
                "Interspeech 2025",
                "https://www.isca-archive.org/interspeech_2025/ohashi25_interspeech.html",
            ),
            paper(
                "2025-09-02",
                "FLM-Audio: Natural Monologues Improves Native Full-Duplex Chatbots via Dual Training",
                "Yiqun Yao et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2509.02521",
            ),
            paper(
                "2025-09-08",
                "FireRedChat: A Pluggable, Full-Duplex Voice Interaction System with Cascaded and Semi-Cascaded Implementations",
                "Junjie Chen et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2509.06502",
            ),
            paper(
                "2025-12",
                "Incorporating Dialogue State Tracking into Japanese Full-duplex Task-oriented Spoken Dialogue Model",
                "Yuya Chiba et al.",
                "IJCNLP-AACL 2025",
                "https://aclanthology.org/2025.findings-ijcnlp.49/",
            ),
            paper(
                "2026-04-30",
                "MiniCPM-o 4.5: Towards Real-Time Full-Duplex Omni-Modal Interaction",
                "Junbo Cui et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2604.27393",
            ),
            paper(
                "2026-05-20",
                "DuplexSLA: A Full-Duplex Spoken Language Model with Synchronized Speech, Language, and Action",
                "Haoyang Zhang et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2605.20755",
            ),
            paper(
                "2026-06-12",
                "BayLing-Duplex: Native Full-Duplex Speech Dialogue with a Single Autoregressive LLM",
                "Qingkai Fang et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2606.14528",
            ),
            paper(
                "2026-07-01",
                "F-Actor: Controllable Conversational Behavior in Full-Duplex Models",
                "Maike Züfle et al.",
                "ACL Findings 2026",
                "https://aclanthology.org/2026.findings-acl.242/",
            ),
            paper(
                "2026-07-02",
                "Hierarchical Acoustic-Semantic Modeling: Modality Separation and Semantic Coherence for Full-Duplex SLMs",
                "Zhenyu Liu et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-long.419/",
            ),
            paper(
                "2026-08-02",
                "JoyAI-Talker: Full-Duplex Speech Interactive Large Model Built for Empathetic Voice Agents",
                "Yinhao Bai et al.",
                "arXiv 2026",
                "https://arxiv.org/abs/2608.01119",
            ),
        ],
    ),
    (
        "Training Methods",
        "Interaction Mode: Turn-taking",
        [
            paper(
                "2023-09",
                "Adapting Text-based Dialogue State Tracker for Spoken Dialogues",
                "Jaeseok Yoon et al.",
                "DSTC 2023",
                "https://aclanthology.org/2023.dstc-1.10/",
            ),
            paper(
                "2025-06-10",
                "Approaching Dialogue State Tracking via Aligning Speech Encoders and LLMs",
                "Šimon Sedláček et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2506.08633",
            ),
            paper(
                "2025-10-10",
                "The Speech-LLM Takes It All: A Truly Fully End-to-End Spoken Dialogue State Tracking Approach",
                "Nizar El Ghazal et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2510.09424",
            ),
            paper(
                "2026-04-16",
                "WavAlign: Enhancing Intelligence and Expressiveness in Spoken Dialogue Models via Adaptive Hybrid Post-Training",
                "Yifu Chen et al.",
                "ACL Findings 2026",
                "https://arxiv.org/abs/2604.14932",
            ),
        ],
    ),
    (
        "Training Methods",
        "Interaction Mode: Half-duplex / Streaming",
        [
            paper(
                "2025-02-19",
                "LLM-Enhanced Dialogue Management for Full-Duplex Spoken Dialogue Systems",
                "Hao Zhang et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2502.14145",
            ),
            paper(
                "2026-07",
                "Dual-Reasoner: Bridging Interleaved Atomicity and Streaming Latency via Thinking-while-Talking",
                "Yangzhuo Li et al.",
                "ACL Findings 2026",
                "https://aclanthology.org/2026.findings-acl.199/",
            ),
        ],
    ),
    (
        "Training Methods",
        "Interaction Mode: Full-duplex",
        [
            paper(
                "2025-07-08",
                "Reinforcement Learning Enhanced Full-Duplex Spoken Dialogue Language Models for Conversational Interactions",
                "Chen Chen et al.",
                "COLM 2025",
                "https://openreview.net/forum?id=QbLbXz8Idp",
            ),
            paper(
                "2025-10-02",
                "Chronological Thinking in Full-Duplex Spoken Dialogue Language Models",
                "Donghang Wu et al.",
                "arXiv 2025",
                "https://arxiv.org/abs/2510.05150",
            ),
            paper(
                "2026-05-01",
                "Reproducing Proficiency-Conditioned Dialogue Features with Full-duplex Spoken Dialogue Models",
                "Takao Obi et al.",
                "IWSDS 2026",
                "https://aclanthology.org/2026.iwsds-1.4/",
            ),
            paper(
                "2026-05-02",
                "Effects of Dialogue Corpora Properties on Fine-Tuning a Moshi-Based Spoken Dialogue Model",
                "Yuto Abe et al.",
                "IWSDS 2026",
                "https://aclanthology.org/2026.iwsds-1.10/",
            ),
            paper(
                "2026-07-01",
                "Dual-Axis Generative Reward Model Toward Semantic and Turn-taking Robustness in Interactive Spoken Dialogue Models",
                "Yifu Chen et al.",
                "ACL 2026",
                "https://aclanthology.org/2026.acl-long.6/",
            ),
            paper(
                "2026-07-02",
                "Sommelier: Scalable Open Multi-turn Audio Pre-processing for Full-duplex Speech Language Models",
                "Kyudan Jung et al.",
                "ACL Industry Track 2026",
                "https://aclanthology.org/2026.acl-industry.18/",
            ),
        ],
    ),
]


def normalize_title(value):
    value = value.lower().replace("‑", "-").replace("–", "-")
    return re.sub(r"[^a-z0-9]+", "", value)


def reference_titles(document):
    result = set()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith('"') and '". ' in text:
            result.add(normalize_title(text.split('". ', 1)[0][1:]))
    return result


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def next_number_id(document, source_id):
    numbering = document.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    source = next(node for node in nums if node.get(qn("w:numId")) == str(source_id))
    new_id = max(int(node.get(qn("w:numId"))) for node in nums) + 1
    clone = deepcopy(source)
    clone.set(qn("w:numId"), str(new_id))
    for override in clone.findall(qn("w:lvlOverride")):
        clone.remove(override)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    clone.append(level_override)
    numbering.append(clone)
    return str(new_id)


def set_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), num_id)
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.insert(0, num_pr)


def add_hyperlink_field(paragraph, url, label="Paper"):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' HYPERLINK "{url}" \\h '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    text_run.extend([run_props, text])
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text_run, end):
        if node.tag == qn("w:r"):
            paragraph._p.append(node)
        else:
            wrapper = OxmlElement("w:r")
            wrapper.append(node)
            paragraph._p.append(wrapper)


def add_paper(document, entry, num_id):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.widow_control = True
    set_numbering(paragraph, num_id)
    paragraph.add_run(f'"{entry["title"]}". ')
    author = paragraph.add_run(f'{entry["author"]} ')
    author.italic = True
    paragraph.add_run(f'{entry["venue"]}. ')
    add_hyperlink_field(paragraph, entry["url"])


def add_heading(document, text, level):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.widow_control = True
    return paragraph


def validate_entries(reference):
    existing = reference_titles(reference)
    seen = set()
    duplicates = []
    collisions = []
    for _, _, entries in SECTIONS:
        dates = [item["date"] for item in entries]
        month_keys = [tuple(int(part) for part in value.split("-")[:2]) for value in dates]
        if month_keys != sorted(month_keys):
            raise ValueError(f"Entries are not chronological: {dates}")
        for item in entries:
            key = normalize_title(item["title"])
            if key in seen:
                duplicates.append(item["title"])
            if key in existing:
                collisions.append(item["title"])
            seen.add(key)
    if duplicates:
        raise ValueError(f"Duplicate entries in supplement: {duplicates}")
    if collisions:
        raise ValueError(f"Entries already present in reference: {collisions}")


def build():
    document = Document(SOURCE)
    validate_entries(document)
    clear_body(document)

    title = document.add_paragraph(style="Title")
    title.add_run("Multi-Turn Speech Dialogue: Supplementary Papers (2022-2026)")
    intro = document.add_paragraph(style="Normal")
    intro.add_run(
        "A supplementary paper list for Multi-Turn Speech Dialogue, searched through "
        "13 August 2026. Papers already listed in readme_restructured.docx are excluded."
    )
    intro.paragraph_format.space_after = Pt(8)

    base_ids = {
        "Survey": 7,
        "Datasets": 9,
        "Evaluation & Benchmarks": 10,
        "Models": 11,
        "Training Methods": 12,
    }
    previous_section = None
    for section, subgroup, entries in SECTIONS:
        if section != previous_section:
            add_heading(document, section, 1)
            previous_section = section
        if subgroup:
            add_heading(document, subgroup, 2)
        num_id = next_number_id(document, base_ids[section])
        for entry in entries:
            add_paper(document, entry, num_id)

    document.core_properties.title = "Multi-Turn Speech Dialogue Supplementary Papers"
    document.core_properties.subject = "Papers from 2022 through 13 August 2026"
    document.core_properties.comments = (
        "Generated as a non-duplicating supplement to readme_restructured.docx."
    )
    document.save(OUTPUT)
    count = sum(len(entries) for _, _, entries in SECTIONS)
    print(f"Created {OUTPUT} with {count} papers")


if __name__ == "__main__":
    build()
