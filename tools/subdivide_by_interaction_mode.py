from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\26547\Desktop\read\readme.docx")
OUTPUT = SOURCE.with_name("readme_interaction_modes.docx")

MODE_HEADINGS = {
    "turn": "Interaction Mode: Turn-taking (\u8f6e\u6d41\u5bf9\u8bdd)",
    "half": "Interaction Mode: Half-duplex / Barge-in (\u534a\u53cc\u5de5 / \u5e26\u6253\u65ad)",
    "full": "Interaction Mode: Full-duplex (\u5168\u53cc\u5de5)",
}

# Papers are assigned only when their main task or architecture establishes the interaction mode.
CLASSIFICATION = {
    "Datasets": {
        "turn": {
            "Advancing Large Language Models to Capture Varied Speaking Styles and Respond Properly in Spoken Conversations (StyleTalk)",
            "Generative Expressive Conversational Speech Synthesis (NCSSD)",
            "DeepDialogue: A Multi-Turn Emotionally-Rich Spoken Dialogue Dataset",
            "UltraVoice: Scaling Fine-Grained Style-Controlled Speech Conversations for Spoken Dialogue Models",
            "Toward Conversational Hungarian Speech Recognition: Introducing the BEA-Large and BEA-Dialogue Datasets",
            "Data-Centric Improvements for Enhancing Multi-Modal Understanding in Spoken Conversation Modeling (ASK-QA)",
            "InteractSpeech: A Speech Dialogue Interaction Corpus for Spoken Dialogue Model",
        },
        "full": {
            "MMedFD: A Real-world Healthcare Benchmark for Multi-turn Full-Duplex Automatic Speech Recognition",
        },
    },
    "Evaluation & Benchmarks": {
        "turn": {
            "SD-Eval: A Benchmark Dataset for Spoken Dialogue Understanding Beyond Words",
            "Contextual Interactive Evaluation of TTS Models in Dialogue Systems",
            "URO-Bench: Towards Comprehensive Evaluation for End-to-End Spoken Dialogue Models",
            "WavReward: Spoken Dialogue Models With Generalist Reward Evaluators",
            "MAD: A Benchmark for Multi-Turn Audio Dialogue Fact-Checking",
            "VoxRole: A Comprehensive Benchmark for Evaluating Speech-Based Role-Playing Agents",
            "Evaluating Bias in Spoken Dialogue LLMs for Real-World Decisions and Recommendations (FairDialogue)",
            "MULTI-Bench: A Multi-turn Interactive Benchmark for Assessing Emotional Intelligence Ability of Spoken Dialogue Models",
            "Audio MultiChallenge: A Multi-Turn Evaluation of Spoken Dialogue Systems on Natural Human Interaction",
            "VoxDialogue: Can Spoken Dialogue Systems Understand Information Beyond Words?",
            "EchoMind: An Interrelated Multi-level Benchmark for Evaluating Empathetic Speech Language Models",
            "HumDial-EIBench: A Human-Recorded Multi-Turn Emotional Intelligence Benchmark for Audio Language Models",
            "VoiceBench: Benchmarking LLM-Based Voice Assistants",
        },
        "half": {
            "Talking Turns: Benchmarking Audio Foundation Models on Turn-Taking Dynamics",
            "Investigating the Impact of Incremental Processing and Voice Activity Projection on Spoken Dialogue Systems",
        },
        "full": {
            "Full-Duplex-Bench: A Benchmark to Evaluate Full-duplex Spoken Dialogue Models on Turn-taking Capabilities",
            "The ICASSP 2026 HumDial Challenge: Benchmarking Human-like Spoken Dialogue Systems in the LLM Era",
        },
    },
    "Models": {
        "turn": {
            "SpeechGPT: Empowering Large Language Models with Intrinsic Cross-Modal Conversational Abilities",
            "Qwen2-Audio Technical Report",
            "Style-Talker: Finetuning Audio Language Model and Style-Based Text-to-Speech Model for Fast Spoken Dialogue Generation",
            "EMOVA: Empowering Language Models to See, Hear and Speak with Vivid Emotions",
            "Internalizing ASR with Implicit Chain of Thought for Efficient Speech-to-Speech Conversational LLM",
            "Building a Taiwanese Mandarin Spoken Language Model: A First Attempt",
            "GLM-4-Voice: Towards Intelligent and Human-Like End-to-End Spoken Chatbot",
            "OpenOmni: Advancing Open-Source Omnimodal Large Language Models with Progressive Multimodal Alignment and Real-time Emotional Speech Synthesis",
            "Qwen2.5-Omni Technical Report",
            "LLaMA-Omni 2: LLM-based Real-time Spoken Chatbot with Autoregressive Streaming Speech Synthesis",
            "GOAT-SLM: A Spoken Language Model with Paralinguistic and Speaker Characteristic Awareness",
            "OpenS2S: Advancing Fully Open-Source End-to-End Empathetic Large Speech Language Model",
            "Step-Audio 2 Technical Report",
            "InteractiveOmni: A Unified Omni-modal Model for Audio-Visual Multi-turn Dialogue",
        },
        "half": {
            "Mini-Omni: Language Models Can Hear, Talk While Thinking in Streaming",
            "Freeze-Omni: A Smart and Low Latency Speech-to-Speech Dialogue Model with Frozen LLM",
        },
        "full": {
            "Beyond Turn-Based Interfaces: Synchronous LLMs as Full-Duplex Dialogue Agents",
            "Moshi: A Speech-Text Foundation Model for Real-Time Dialogue",
            "SALMONN-omni: A Codec-free LLM for Full-duplex Speech Understanding and Generation",
            "MinMo: A Multimodal Large Language Model for Seamless Voice Interaction",
            "SALM-Duplex: Efficient and Direct Duplex Modeling for Speech-to-Speech Language Model",
            "OmniFlatten: An End-to-end GPT Model for Seamless Voice Conversation",
            "FlashLabs Chroma 1.0: A Real-Time End-to-End Spoken Dialogue Model with Personalized Voice Cloning",
            "PersonaPlex: Voice and Role Control for Full Duplex Conversational Speech Models",
        },
    },
    "Training Methods": {
        "turn": {
            "SpeechAlign: Aligning Speech Generation to Human Preferences",
            "WavRAG: Audio-Integrated Retrieval Augmented Generation for Spoken Dialogue Models",
            "Enhancing Speech-to-Speech Dialogue Modeling with End-to-End Retrieval-Augmented Generation",
            "Aligning Spoken Dialogue Models from User Interactions",
            "Chain-of-Thought Training for Open E2E Spoken Dialogue Systems",
            "Stream RAG: Instant and Accurate Spoken Dialogue Systems with Streaming Tool Usage",
            "VoiceTextBlender: Augmenting Large Language Models with Speech Capabilities via Single-Stage Joint Speech-Text Supervised Fine-Tuning",
            "Optimizing Conversational Quality in Spoken Dialogue Systems with Reinforcement Learning from AI Feedback",
        },
        "full": {
            "SHANKS: Simultaneous Hearing and Thinking for Spoken Language Models",
            "Multi-Faceted Interactivity Alignment in Full-Duplex Speech Models",
        },
    },
}


def paper_title(paragraph):
    text = paragraph.text.strip()
    if not text.startswith('"'):
        return None
    return text.split('". ', 1)[0][1:]


def section_papers(document, heading):
    capture = False
    result = []
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 1":
            if capture:
                break
            capture = paragraph.text.strip() == heading
            continue
        if capture and paper_title(paragraph):
            result.append(paragraph)
    return result


def fresh_numbering_id(document, source_id):
    numbering = document.part.numbering_part.element
    source = next(node for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId")) == source_id)
    values = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = str(max(values) + 1)
    clone = deepcopy(source)
    clone.set(qn("w:numId"), new_id)
    numbering.append(clone)
    return new_id


def apply_num_id(paragraph, num_id):
    properties = paragraph._p.get_or_add_pPr()
    number_properties = properties.find(qn("w:numPr"))
    if number_properties is None:
        number_properties = OxmlElement("w:numPr")
        properties.append(number_properties)
    for child in list(number_properties):
        if child.tag in {qn("w:ilvl"), qn("w:numId")}:
            number_properties.remove(child)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), num_id)
    number_properties.append(level)
    number_properties.append(number)


def new_heading_after(document, cursor, text):
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.add_run(text)
    cursor.addnext(paragraph._p)
    return paragraph._p


def main():
    document = Document(SOURCE)
    for section, modes in CLASSIFICATION.items():
        heading = next(p for p in document.paragraphs if p.style.name == "Heading 1" and p.text.strip() == section)
        papers = section_papers(document, section)
        by_title = {paper_title(paragraph): paragraph for paragraph in papers}
        classified = set().union(*modes.values())
        if set(by_title) != classified:
            missing = set(by_title) - classified
            unknown = classified - set(by_title)
            raise ValueError(f"{section}: missing={missing}, unknown={unknown}")

        cursor = heading._p
        for mode in ("turn", "half", "full"):
            titles = modes.get(mode, set())
            if not titles:
                continue
            cursor = new_heading_after(document, cursor, MODE_HEADINGS[mode])
            source_num = by_title[next(iter(titles))]._p.pPr.find(qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
            num_id = fresh_numbering_id(document, source_num)
            for paragraph in papers:
                if paper_title(paragraph) not in titles:
                    continue
                cursor.addnext(paragraph._p)
                apply_num_id(paragraph, num_id)
                cursor = paragraph._p
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
