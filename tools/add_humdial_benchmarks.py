from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "readme.docx"
OUTPUT = ROOT / "readme_humdial_updated.docx"

PAPERS = [
    {
        "anchor": "VoxDialogue: Can Spoken Dialogue Systems Understand Information Beyond Words?",
        "text": '"EchoMind: An Interrelated Multi-level Benchmark for Evaluating Empathetic Speech Language Models". Li Zhou et al. arXiv 2025. ',
        "url": "https://arxiv.org/abs/2510.22758",
    },
    {
        "anchor": "The ICASSP 2026 HumDial Challenge: Benchmarking Human-like Spoken Dialogue Systems in the LLM Era",
        "text": '"HumDial-EIBench: A Human-Recorded Multi-Turn Emotional Intelligence Benchmark for Audio Language Models". Shuiyuan Wang et al. arXiv 2026. ',
        "url": "https://arxiv.org/abs/2604.11594",
    },
]


def add_hyperlink(paragraph, label, url):
    hyperlink = OxmlElement("w:hyperlink")
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_after(anchor, text, url):
    new_paragraph = anchor.insert_paragraph_before()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    anchor._p.addnext(new_paragraph._p)
    new_paragraph.style = anchor.style
    new_paragraph.paragraph_format._element = deepcopy(anchor.paragraph_format._element)
    new_paragraph.add_run(text)
    add_hyperlink(new_paragraph, "Paper", url)


def main():
    document = Document(SOURCE)
    existing = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for paper in PAPERS:
        if paper["text"].split('". ')[0].removeprefix('"') in existing:
            raise ValueError("Paper is already present: " + paper["text"])
        anchor = next(
            paragraph for paragraph in document.paragraphs if paper["anchor"] in paragraph.text
        )
        insert_after(anchor, paper["text"], paper["url"])
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
