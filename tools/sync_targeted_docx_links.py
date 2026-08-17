from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
README = ROOT / "README.md"
DOCX = ROOT / "readme.docx"
OUTPUT = ROOT / "readme_updated.docx"

TARGET_TITLES = {
    "DeepDialogue: A Multi-Turn Emotionally-Rich Spoken Dialogue Dataset",
    "The ICASSP 2026 HumDial Challenge: Benchmarking Human-like Spoken Dialogue Systems in the LLM Era",
    "InteractiveOmni: A Unified Omni-modal Model for Audio-Visual Multi-turn Dialogue",
    "SpeechGPT: Empowering Large Language Models with Intrinsic Cross-Modal Conversational Abilities",
    "PersonaPlex: Voice and Role Control for Full Duplex Conversational Speech Models",
    "Enhancing Speech-to-Speech Dialogue Modeling with End-to-End Retrieval-Augmented Generation",
    "Multi-Faceted Interactivity Alignment in Full-Duplex Speech Models",
}


def read_links():
    pattern = re.compile(r'^\d+\. \*\*"(?P<title>.+?)"\*\*\..*?\[\[Paper\]\((?P<url>https://[^)]+)\)\]$')
    return {
        match.group("title"): match.group("url")
        for line in README.read_text(encoding="utf-8").splitlines()
        if (match := pattern.match(line))
    }


def append_hyperlink(paragraph, text, url):
    hyperlink = OxmlElement("w:hyperlink")
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def main():
    links = read_links()
    doc = Document(DOCX)
    updated = 0
    for paragraph in doc.paragraphs:
        title_match = re.search(r'"(.+?)"\. ', paragraph.text.strip())
        if not title_match:
            continue
        title = title_match.group(1)
        if title == "VoiceBench: Benchmarking LLM-Based Voice Assistants":
            paragraph._element.getparent().remove(paragraph._element)
            continue
        if title not in TARGET_TITLES:
            continue
        hyperlink_nodes = paragraph._p.xpath('.//w:hyperlink')
        if hyperlink_nodes:
            continue
        paragraph.add_run(" ")
        append_hyperlink(paragraph, "Paper", links[title])
        updated += 1
    if updated != len(TARGET_TITLES):
        raise ValueError(f"Expected {len(TARGET_TITLES)} updates, applied {updated}")
    doc.save(OUTPUT)
    print(f"Added {updated} links and removed VoiceBench in {OUTPUT}")


if __name__ == "__main__":
    main()
