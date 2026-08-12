from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "readme_humdial_numbered.docx"
OUTPUT = ROOT / "readme_survey_numbered.docx"

SURVEYS = [
    {
        "anchor": "On The Landscape of Spoken Language Models: A Comprehensive Survey",
        "text": '"From Turn-Taking to Synchronous Dialogue: A Survey of Full-Duplex Spoken Language Models". Yuxuan Chen and Haoyuan Yu. arXiv 2025. ',
        "url": "https://arxiv.org/abs/2509.14515",
    },
    {
        "anchor": "Recent Advances in Speech Language Models: A Survey",
        "text": '"A Survey of Full-Duplex Spoken Dialogue Systems: Architectural Hierarchy, Interaction Ontology, and Decision State Machine". Jingyu Lu et al. arXiv 2026. ',
        "url": "https://arxiv.org/abs/2606.19453",
    },
]


def add_hyperlink(paragraph, label, url):
    hyperlink = OxmlElement("w:hyperlink")
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_after(anchor, text, url):
    paragraph = anchor.insert_paragraph_before()
    paragraph._p.getparent().remove(paragraph._p)
    anchor._p.addnext(paragraph._p)
    paragraph.style = anchor.style
    paragraph.paragraph_format._element = deepcopy(anchor.paragraph_format._element)
    paragraph.add_run(text)
    add_hyperlink(paragraph, "Paper", url)


def remove_number_prefix(paragraph):
    if paragraph.runs:
        paragraph.runs[0].text = re.sub(r"^\d+\.\s*", "", paragraph.runs[0].text)


def add_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    format_node = OxmlElement("w:numFmt")
    format_node.set(qn("w:val"), "decimal")
    level.append(format_node)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_auto_numbering(document):
    current_num_id = None
    numbered = 0
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 1":
            current_num_id = add_numbering(document)
            continue
        if current_num_id is None or not paragraph.text.rstrip().endswith("Paper"):
            continue
        remove_number_prefix(paragraph)
        properties = paragraph._p.get_or_add_pPr()
        num_properties = properties.find(qn("w:numPr"))
        if num_properties is None:
            num_properties = OxmlElement("w:numPr")
            properties.append(num_properties)
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(current_num_id))
        for child in list(num_properties):
            if child.tag in {qn("w:ilvl"), qn("w:numId")}:
                num_properties.remove(child)
        num_properties.append(level)
        num_properties.append(num)
        numbered += 1
    return numbered


def main():
    document = Document(SOURCE)
    existing = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for survey in SURVEYS:
        title = survey["text"].split('". ')[0].removeprefix('"')
        if title in existing:
            raise ValueError(f"Survey already present: {title}")
        anchor = next(p for p in document.paragraphs if survey["anchor"] in p.text)
        insert_after(anchor, survey["text"], survey["url"])

    numbered = apply_auto_numbering(document)
    document.save(OUTPUT)
    print(f"Created {OUTPUT} with {numbered} automatically numbered entries")


if __name__ == "__main__":
    main()
